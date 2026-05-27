from __future__ import annotations

import os
from datetime import date, timedelta
from io import BytesIO
from pathlib import Path
from typing import Any

import streamlit as st
from docx import Document
from dotenv import load_dotenv
from langchain_anthropic import ChatAnthropic
from langchain_core.messages import HumanMessage, SystemMessage
from pydantic import BaseModel, Field

from db import PatientRow, SessionRow, assemble_transcript, fetch_patients, fetch_sessions


BASE_DIR = Path(__file__).parent
FORMAT_DIR = BASE_DIR / "format"
PROMPTS_DIR = BASE_DIR / "prompts"
TYPE_PROMPTS_DIR = PROMPTS_DIR / "type_prompts"
CORE_RULES_V3 = PROMPTS_DIR / "core_rules_v3.txt"
DYNAMIC_EXCLUDED_SESSION_TYPES = {"Room Change", "Section Change"}
DEFAULT_MODEL = "claude-sonnet-4-6"
DEFAULT_MAX_OUTPUT_TOKENS = 64000

# v1.0/v2.0 are single-file system prompts (+ a shared dynamic-sections file).
# v3.0 is a composed prompt: one shared core_rules file + one focused module per session type.
PROMPT_VERSION_LABELS = ["v1.0", "v2.0", "v3.0"]

PROMPT_VERSIONS: dict[str, Path] = {
    "v1.0": PROMPTS_DIR / "lumio_system_prompt_v1.txt",
    "v2.0": PROMPTS_DIR / "lumio_system_prompt_v2.txt",
}

DYNAMIC_PROMPT_VERSIONS: dict[str, Path] = {
    "v1.0": PROMPTS_DIR / "lumio_dynamic_sections_3_4_prompt_v1.txt",
    "v2.0": PROMPTS_DIR / "lumio_dynamic_sections_3_4_prompt_v2.txt",
}


def type_prompt_path(session_type: str) -> Path:
    """Map a session-type display name to its v3.0 module file (spaces → underscores)."""
    return TYPE_PROMPTS_DIR / f"{session_type.replace(' ', '_')}.txt"


class ClinicalSummaryOutput(BaseModel):
    """Validated response envelope for the generated clinical summary."""

    rendered_summary: str = Field(
        description=(
            "The completed clinical summary, rendered in the same Markdown headings, "
            "field order, tables, checkboxes, and visible structure as the selected "
            "session template. Return the full note from the first heading to the last field."
        )
    )


def read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8").strip()


def read_docx_text(path_or_file: Path | Any) -> str:
    document = Document(path_or_file)
    blocks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    return "\n".join(blocks).strip()


def discover_session_templates() -> dict[str, Path]:
    templates = {
        path.stem.replace("_", " "): path
        for path in sorted(FORMAT_DIR.glob("*.md"), key=lambda item: item.stem.lower())
    }

    if templates:
        return templates

    return {
        path.stem.replace("_", " "): path
        for path in sorted(FORMAT_DIR.glob("*.docx"), key=lambda item: item.stem.lower())
    }


def read_template_text(path: Path) -> str:
    if path.suffix.lower() == ".md":
        return read_text_file(path)

    if path.suffix.lower() == ".docx":
        return read_docx_text(path)

    raise ValueError(f"Unsupported template format: {path.suffix}")


def should_include_dynamic_prompt(session_type: str) -> bool:
    return session_type not in DYNAMIC_EXCLUDED_SESSION_TYPES


OUTPUT_REQUIREMENTS = """
<output_requirements>
These rules govern Markdown structure and formatting only. All clinical content rules are defined in the system prompt above and must be followed as written.

HEADINGS: Use the exact heading levels (#, ##, ###) and heading text from the selected template. Do not add or rename headings. Do not remove a heading that has content. (If the clinical content rules above require omitting an empty section, follow those rules.)

FIELD FORMAT: Each field and its value on its own line. Mirror the list style of the template exactly (e.g. `- **Field:**`). Never merge multiple fields onto one line.

TABLES: Reproduce Markdown tables as consecutive lines with NO indentation — every line starts at the left margin. The separator row (|---|---|...|) must come immediately after the header row, with NO blank line between header, separator, and data rows. Put exactly one blank line before the header row and one blank line after the last data row. Never collapse rows onto a single line, and never indent table lines (indented tables render as plain text).
Example:

| Medication | Dose | Frequency | Duration | Purpose |
|---|---|---|---|---|
| Sertraline | 50 mg | Morning | 5 days | Depression |

CHECKBOXES: Mark the transcript-supported option with `☑` or `- [x]`. Leave others as `☐` or `- [ ]`.

COMPLETENESS: Never stop mid-section or mid-table. Output every field and section that has content, through to the end of the note. (The clinical content rules above govern which fields and sections are included.)
</output_requirements>
""".strip()


def build_system_prompt(session_type: str, version: str = "v1.0") -> str:
    if version == "v3.0":
        core = read_text_file(CORE_RULES_V3)
        module_path = type_prompt_path(session_type)
        if module_path.exists():
            return f"{core}\n\n{read_text_file(module_path)}\n\n{OUTPUT_REQUIREMENTS}"
        return f"{core}\n\n{OUTPUT_REQUIREMENTS}"

    system_prompt = read_text_file(PROMPT_VERSIONS[version])

    include_dynamic = version == "v2.0" or should_include_dynamic_prompt(session_type)
    if include_dynamic:
        dynamic_prompt = read_text_file(DYNAMIC_PROMPT_VERSIONS[version])
        return f"{system_prompt}\n\n{dynamic_prompt}\n\n{OUTPUT_REQUIREMENTS}"

    return f"{system_prompt}\n\n{OUTPUT_REQUIREMENTS}"


def build_user_prompt(session_type: str, template_text: str, transcript: str) -> str:
    return f"""
<task>
Generate one complete clinical summary for the selected session type.
</task>

<selected_session_type>
{session_type}
</selected_session_type>

<selected_session_format>
{template_text}
</selected_session_format>

<verified_session_transcript>
{transcript}
</verified_session_transcript>
""".strip()


def generate_summary(
    *,
    session_type: str,
    template_text: str,
    transcript: str,
    model_name: str,
    max_output_tokens: int,
    prompt_version: str = "v1.0",
) -> ClinicalSummaryOutput:
    model = ChatAnthropic(
        model=model_name,
        max_tokens=max_output_tokens,
        timeout=240,
        max_retries=2,
        temperature=0,
    )
    structured_model = model.with_structured_output(ClinicalSummaryOutput)

    response = structured_model.invoke(
        [
            SystemMessage(content=build_system_prompt(session_type, prompt_version)),
            HumanMessage(content=build_user_prompt(session_type, template_text, transcript)),
        ]
    )

    if not isinstance(response, ClinicalSummaryOutput):
        return ClinicalSummaryOutput.model_validate(response)

    return response


def get_uploaded_text(uploaded_file: Any) -> str:
    if uploaded_file.name.lower().endswith(".docx"):
        return read_docx_text(uploaded_file)

    return uploaded_file.getvalue().decode("utf-8")


def is_table_row(line: str) -> bool:
    return " | " in line and len([cell for cell in line.split("|") if cell.strip()]) > 1


def add_text_paragraph(document: Document, line: str) -> None:
    paragraph = document.add_paragraph()
    stripped = line.strip()

    if ":" in stripped:
        label, value = stripped.split(":", 1)
        paragraph.add_run(f"{label}:").bold = True
        if value:
            paragraph.add_run(value)
        return

    run = paragraph.add_run(stripped)
    if stripped.isupper() or stripped.endswith("NOTE") or stripped.endswith("TEMPLATE"):
        run.bold = True


def add_table(document: Document, rows: list[str]) -> None:
    parsed_rows = [[cell.strip() for cell in row.split("|")] for row in rows]
    max_columns = max(len(row) for row in parsed_rows)
    table = document.add_table(rows=len(parsed_rows), cols=max_columns)
    table.style = "Table Grid"

    for row_index, row in enumerate(parsed_rows):
        for col_index in range(max_columns):
            cell_text = row[col_index] if col_index < len(row) else ""
            cell = table.cell(row_index, col_index)
            cell.text = cell_text
            if row_index == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def render_summary_docx(summary: str) -> bytes:
    document = Document()
    pending_table_rows: list[str] = []

    def flush_table() -> None:
        if pending_table_rows:
            add_table(document, pending_table_rows)
            pending_table_rows.clear()

    for raw_line in summary.splitlines():
        line = raw_line.strip()
        if not line:
            flush_table()
            document.add_paragraph()
            continue

        if is_table_row(line):
            pending_table_rows.append(line)
            continue

        flush_table()
        add_text_paragraph(document, line)

    flush_table()

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _static_context(patient: PatientRow, session: SessionRow) -> str:
    """Pre-fill the known static fields so the LLM doesn't have to guess them.

    Only fields with an actual value are included — empty fields are omitted so
    the model does not echo a blank label into the output.
    """
    fields = [
        ("Date", session.session_date.strftime("%d %B %Y") if session.session_date else ""),
        ("Session Number", str(session.session_number) if session.session_number else ""),
        ("Mode of Consultation", session.modality_label),
        ("Clinician", session.clinician),
        ("Patient Name", f"{patient.first_name} {patient.last_name}".strip()),
        ("Age / Gender", patient.age_gender),
    ]
    return "\n".join(f"{label}: {value}" for label, value in fields if value)


def show_login() -> None:
    st.title("Lumio Clinical Summary")
    st.subheader("Sign in")

    with st.form("login_form"):
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")
        submitted = st.form_submit_button("Login", use_container_width=True)

    if submitted:
        expected_user = os.environ.get("APP_USERNAME", "")
        expected_pass = os.environ.get("APP_PASSWORD", "")
        if username == expected_user and password == expected_pass:
            st.session_state["authenticated"] = True
            st.rerun()
        else:
            st.error("Invalid username or password.")


def main() -> None:
    load_dotenv()

    st.set_page_config(page_title="Lumio Clinical Summary", page_icon="📝", layout="wide")

    if not st.session_state.get("authenticated"):
        show_login()
        st.stop()

    st.title("Lumio Clinical Summary Generator")

    templates = discover_session_templates()
    if not templates:
        st.error("No .md or .docx templates found in the format folder.")
        st.stop()

    MODEL_OPTIONS = {
        "Claude Sonnet 4.6": "claude-sonnet-4-6",
        "Claude Opus 4.6": "claude-opus-4-6",
    }

    # ── Sidebar ──────────────────────────────────────────────────────────────
    with st.sidebar:
        st.header("Settings")
        selected_model_label = st.selectbox("Model", options=list(MODEL_OPTIONS.keys()))
        model_name = MODEL_OPTIONS[selected_model_label]
        prompt_version = st.selectbox("Prompt Version", options=PROMPT_VERSION_LABELS, index=len(PROMPT_VERSION_LABELS) - 1)

    # ── Step 1: Patient selection ─────────────────────────────────────────────
    try:
        patients: list[PatientRow] = fetch_patients()
    except Exception as exc:
        st.error(f"Database connection failed: {exc}")
        st.stop()

    if not patients:
        st.warning("No patients with available transcripts found.")
        st.stop()

    DATE_FILTER_OPTIONS: dict[str, Any] = {
        "All": None,
        "Today": "today",
        "Yesterday": "yesterday",
        "Last 7 days": timedelta(days=7),
        "Last 30 days": timedelta(days=30),
        "Last 3 months": timedelta(days=90),
        "Last 6 months": timedelta(days=180),
        "Last year": timedelta(days=365),
    }

    f_col1, f_col2, f_col3 = st.columns([3, 2, 1])
    with f_col1:
        name_search = st.text_input("Search patient", placeholder="Type name or ID…", label_visibility="collapsed")
    with f_col2:
        date_filter_label = st.selectbox("Last session", options=list(DATE_FILTER_OPTIONS.keys()), label_visibility="collapsed")
    with f_col3:
        sort_order = st.selectbox("Sort", options=["A → Z", "Z → A"], label_visibility="collapsed")

    filtered = patients

    if name_search:
        # Match every whitespace-separated token against first + last + ID so
        # "Jane Doe" works (previously the full string had to appear in one field).
        tokens = [t for t in name_search.strip().lower().split() if t]
        if tokens:
            filtered = [
                p
                for p in filtered
                if all(
                    tok in f"{p.first_name} {p.last_name} {p.custom_patient_id}".lower()
                    for tok in tokens
                )
            ]

    cutoff = DATE_FILTER_OPTIONS[date_filter_label]
    if cutoff == "today":
        filtered = [p for p in filtered if p.last_session_date == date.today()]
    elif cutoff == "yesterday":
        yesterday = date.today() - timedelta(days=1)
        filtered = [p for p in filtered if p.last_session_date == yesterday]
    elif cutoff is not None:
        threshold = date.today() - cutoff
        filtered = [p for p in filtered if p.last_session_date and p.last_session_date >= threshold]

    filtered = sorted(
        filtered,
        key=lambda p: (p.first_name.lower(), p.last_name.lower()),
        reverse=(sort_order == "Z → A"),
    )

    if not filtered:
        st.info("No patients match the current filters.")
        st.stop()

    patient_map: dict[str, PatientRow] = {p.display_name: p for p in filtered}
    selected_patient_name = st.selectbox("Select Patient", options=list(patient_map.keys()))
    patient = patient_map[selected_patient_name]

    # ── Step 2: Session selection ─────────────────────────────────────────────
    sessions: list[SessionRow] = fetch_sessions(patient.clientid)

    if not sessions:
        st.info("No sessions with transcripts found for this patient.")
        st.stop()

    session_map: dict[str, SessionRow] = {s.display_label: s for s in sessions}
    selected_session_label = st.selectbox("Select Session", options=list(session_map.keys()))
    session = session_map[selected_session_label]

    # ── Step 3: Session preview ───────────────────────────────────────────────
    with st.expander("Session Preview", expanded=True):
        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown(f"**Patient:** {patient.first_name} {patient.last_name}")
            st.markdown(f"**Age / Gender:** {patient.age_gender or '—'}")
            st.markdown(f"**Patient ID:** {patient.custom_patient_id}")
        with col2:
            st.markdown(f"**Date:** {session.session_date.strftime('%d %B %Y')}")
            st.markdown(f"**Session #:** {session.session_number}")
            st.markdown(f"**Type:** {session.session_type.replace('_', ' ').title()}")
        with col3:
            st.markdown(f"**Clinician:** {session.clinician}")
            st.markdown(f"**Modality:** {session.modality_label}")
            st.markdown(f"**Status:** {session.status.replace('_', ' ').title()}")

    # ── Step 3b: Transcript preview ──────────────────────────────────────────
    with st.expander("Transcript Preview", expanded=False):
        show_translation = st.toggle("Show translation", value=True)
        transcript_preview = assemble_transcript(session, use_translation=show_translation)
        st.text_area(
            label="",
            value=transcript_preview,
            height=320,
            disabled=True,
            label_visibility="collapsed",
        )

    # ── Step 4: Note type override ────────────────────────────────────────────
    template_options = list(templates.keys())
    suggested = session.suggested_template
    default_index = template_options.index(suggested) if suggested in template_options else 0

    session_type = st.selectbox(
        "Note Type",
        options=template_options,
        index=default_index,
        help="Auto-selected based on session type. Override if needed.",
    )

    if prompt_version == "v3.0":
        if type_prompt_path(session_type).exists():
            st.caption(f"Core rules + “{session_type}” module (type-specific fields & vocabulary).")
        else:
            st.caption("Core rules only (no type-specific module found for this note type).")
    else:
        include_dynamic = prompt_version == "v2.0" or should_include_dynamic_prompt(session_type)
        if include_dynamic:
            st.caption("Dynamic sections 3 & 4 (controlled vocabularies + screening checklists) included.")
        else:
            st.caption("Dynamic sections 3 & 4 excluded for this note type.")

    with st.expander("System Prompt", expanded=False):
        st.text(build_system_prompt(session_type, prompt_version))

    # ── Step 5: Generate ─────────────────────────────────────────────────────
    if st.button("Generate Summary", type="primary"):
        template_path = templates[session_type]
        template_text = read_template_text(template_path)
        transcript_text = assemble_transcript(session)
        static_ctx = _static_context(patient, session)

        # Inject static fields into the transcript block so the LLM has them
        full_context = f"<session_metadata>\n{static_ctx}\n</session_metadata>\n\n{transcript_text}"

        with st.spinner("Generating clinical summary…"):
            try:
                output = generate_summary(
                    session_type=session_type,
                    template_text=template_text,
                    transcript=full_context,
                    model_name=model_name,
                    max_output_tokens=DEFAULT_MAX_OUTPUT_TOKENS,
                    prompt_version=prompt_version,
                )
            except Exception as exc:
                st.error(f"Generation failed: {exc}")
                st.stop()

        st.subheader("Generated Summary")
        st.markdown(output.rendered_summary)
        st.download_button(
            "Download as .docx",
            data=render_summary_docx(output.rendered_summary),
            file_name=f"{patient.last_name}_{session.session_date}_session{session.session_number}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


if __name__ == "__main__":
    main()
